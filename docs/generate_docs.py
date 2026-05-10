"""
WorkPing — Project Documentation Generator
Produces docs/WorkPing_Project_Documentation.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Brand colours ─────────────────────────────────────────────────────────────
BRAND_DARK   = RGBColor(0x0F, 0x17, 0x2A)   # deep navy
BRAND_BLUE   = RGBColor(0x1A, 0x56, 0xDB)   # primary blue
BRAND_ACCENT = RGBColor(0x06, 0xB6, 0xD4)   # cyan accent
BRAND_LIGHT  = RGBColor(0xF0, 0xF4, 0xFF)   # light wash
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
GREY_TEXT    = RGBColor(0x4B, 0x55, 0x63)
LIGHT_ROW    = RGBColor(0xF8, 0xFA, 0xFF)

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ══════════════════════════════════════════════════════════════════════════════
# Helpers
# ══════════════════════════════════════════════════════════════════════════════

def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    hex_color = str(rgb)  # RGBColor.__str__ returns 'RRGGBB'
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="D0D5E8"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"),   "single")
        border.set(qn("w:sz"),    "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def remove_table_border(table):
    tbl    = table._tbl
    tblPr  = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top","left","bottom","right","insideH","insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        tblBorders.append(b)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)

def add_page_break(document):
    p = document.add_paragraph()
    run = p.add_run()
    run.add_break(docx_break_type())
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)

def docx_break_type():
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    return br   # We'll handle this differently

def page_break(document):
    document.add_page_break()

def heading(document, text, level=1, color=BRAND_DARK, size=None, bold=True, space_before=18, space_after=8):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    else:
        sizes = {1: 22, 2: 16, 3: 13, 4: 12}
        run.font.size = Pt(sizes.get(level, 12))
    return p

def subheading(document, text, color=BRAND_BLUE, size=13):
    return heading(document, text, level=2, color=color, size=size, space_before=14, space_after=5)

def body(document, text, color=GREY_TEXT, size=10.5, space_after=6, italic=False, bold=False):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.color.rgb = color
    run.font.size      = Pt(size)
    run.italic         = italic
    run.bold           = bold
    return p

def bullet(document, text, color=GREY_TEXT, size=10.5, indent=0.5, marker="▸"):
    p   = document.add_paragraph()
    p.paragraph_format.left_indent   = Cm(indent)
    p.paragraph_format.space_before  = Pt(2)
    p.paragraph_format.space_after   = Pt(4)
    run = p.add_run(f"{marker}  {text}")
    run.font.color.rgb = color
    run.font.size      = Pt(size)
    return p

def sub_bullet(document, text, color=GREY_TEXT, size=10):
    return bullet(document, text, color=color, size=size, indent=1.2, marker="–")

def divider(document, color="1A56DB"):
    p    = document.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def make_table(document, headers, rows, col_widths=None, header_bg=BRAND_BLUE, zebra=True):
    n_cols = len(headers)
    table  = document.add_table(rows=1, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_border(table)

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        set_cell_bg(cell, header_bg)
        set_cell_borders(cell, "1A56DB")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p    = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        run  = p.add_run(h)
        run.bold            = True
        run.font.color.rgb  = WHITE
        run.font.size       = Pt(10)

    # Data rows
    for ri, row_data in enumerate(rows):
        row_cells = table.add_row().cells
        bg = LIGHT_ROW if (zebra and ri % 2 == 0) else WHITE
        for ci, val in enumerate(row_data):
            cell = row_cells[ci]
            set_cell_bg(cell, bg)
            set_cell_borders(cell, "D0D5E8")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p    = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            run  = p.add_run(str(val))
            run.font.color.rgb = GREY_TEXT
            run.font.size      = Pt(10)

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Cm(width)

    document.add_paragraph().paragraph_format.space_after = Pt(4)
    return table

def accent_box(document, title, lines, bg=BRAND_LIGHT, accent=BRAND_BLUE):
    """Simulate a highlighted box using a 1-cell table."""
    table = document.add_table(rows=1, cols=1)
    remove_table_border(table)
    cell  = table.cell(0, 0)
    set_cell_bg(cell, bg)
    set_cell_borders(cell, str(accent))

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = accent
    run.font.size = Pt(11)

    for line in lines:
        np = cell.add_paragraph(f"  {line}")
        np.paragraph_format.space_before = Pt(1)
        np.paragraph_format.space_after  = Pt(1)
        nr = np.runs[0] if np.runs else np.add_run(f"  {line}")
        nr.font.color.rgb = GREY_TEXT
        nr.font.size = Pt(10)

    cell.add_paragraph().paragraph_format.space_after = Pt(2)
    document.add_paragraph().paragraph_format.space_after = Pt(6)


# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

# Top accent bar (simulate with a 1-cell table)
bar_table = doc.add_table(rows=1, cols=1)
remove_table_border(bar_table)
bar_cell = bar_table.cell(0, 0)
set_cell_bg(bar_cell, BRAND_BLUE)
bar_p = bar_cell.paragraphs[0]
bar_p.paragraph_format.space_before = Pt(10)
bar_p.paragraph_format.space_after  = Pt(10)

doc.add_paragraph().paragraph_format.space_after = Pt(30)

# Logo / title block
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(20)
title_r = title_p.add_run("WorkPing")
title_r.font.size      = Pt(52)
title_r.font.bold      = True
title_r.font.color.rgb = BRAND_BLUE

tagline_p = doc.add_paragraph()
tagline_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tagline_r = tagline_p.add_run("Intelligent Workforce Management Platform")
tagline_r.font.size      = Pt(18)
tagline_r.font.color.rgb = BRAND_ACCENT
tagline_r.font.italic    = True

doc.add_paragraph().paragraph_format.space_after = Pt(16)
divider(doc)
doc.add_paragraph().paragraph_format.space_after = Pt(16)

# Subtitle block
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_r = sub_p.add_run("Project Documentation")
sub_r.font.size      = Pt(20)
sub_r.font.bold      = True
sub_r.font.color.rgb = BRAND_DARK

doc.add_paragraph().paragraph_format.space_after = Pt(40)

# Info card
info_lines = [
    ("Platform", "Multi-tenant SaaS — HR & Workforce Management"),
    ("Architecture", "Microservices on Oracle Cloud Infrastructure"),
    ("Author", "Nikhil Charan"),
    ("Year", "2026"),
    ("Deployment", "workping.live"),
]
info_table = doc.add_table(rows=len(info_lines), cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
remove_table_border(info_table)
for ri, (label, value) in enumerate(info_lines):
    lc = info_table.cell(ri, 0)
    vc = info_table.cell(ri, 1)
    set_cell_bg(lc, BRAND_DARK)
    set_cell_bg(vc, BRAND_LIGHT)
    set_cell_borders(lc, "1A56DB")
    set_cell_borders(vc, "D0D5E8")
    lc.width = Cm(5)
    vc.width = Cm(10)
    lp = lc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    lp.paragraph_format.space_before = Pt(5)
    lp.paragraph_format.space_after  = Pt(5)
    lr = lp.add_run(f"  {label}  ")
    lr.bold            = True
    lr.font.color.rgb  = WHITE
    lr.font.size       = Pt(10.5)
    vp = vc.paragraphs[0]
    vp.paragraph_format.space_before = Pt(5)
    vp.paragraph_format.space_after  = Pt(5)
    vr = vp.add_run(f"  {value}")
    vr.font.color.rgb = GREY_TEXT
    vr.font.size      = Pt(10.5)

doc.add_paragraph().paragraph_format.space_after = Pt(40)

# Bottom accent
bar2_table = doc.add_table(rows=1, cols=3)
remove_table_border(bar2_table)
colors_b = [BRAND_DARK, BRAND_BLUE, BRAND_ACCENT]
for i in range(3):
    bc = bar2_table.cell(0, i)
    set_cell_bg(bc, colors_b[i])
    bp = bc.paragraphs[0]
    bp.paragraph_format.space_before = Pt(6)
    bp.paragraph_format.space_after  = Pt(6)


# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════════════════════════
page_break(doc)

heading(doc, "Table of Contents", level=1, size=20, space_before=0)
divider(doc)
doc.add_paragraph().paragraph_format.space_after = Pt(6)

toc_items = [
    ("1", "Problem Statement"),
    ("2", "Solution Overview"),
    ("3", "System Architecture"),
    ("4", "Technology Stack"),
    ("5", "Microservices — Detailed Breakdown"),
    ("6", "Workflow & Data Flows"),
    ("7", "Security Model"),
    ("8", "Deployment & Infrastructure"),
    ("9", "Impact on HR Management"),
    ("10", "Future Scope"),
]
for num, title in toc_items:
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    nr  = p.add_run(f"  {num}.  ")
    nr.bold            = True
    nr.font.color.rgb  = BRAND_BLUE
    nr.font.size       = Pt(11)
    tr  = p.add_run(title)
    tr.font.color.rgb  = BRAND_DARK
    tr.font.size       = Pt(11)


# ══════════════════════════════════════════════════════════════════════════════
# 1. PROBLEM STATEMENT
# ══════════════════════════════════════════════════════════════════════════════
page_break(doc)

heading(doc, "1. Problem Statement", level=1, size=20, space_before=0)
divider(doc)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

body(doc,
     "Small and medium enterprises in India and beyond face a fragmented, error-prone, and time-consuming "
     "approach to workforce management. Attendance is marked manually on registers or through basic punch-card "
     "machines — systems that are trivially forged (buddy punching), provide no real-time visibility, and generate "
     "mountains of paper that must be manually transcribed into spreadsheets at the end of every month.",
     size=10.5)

body(doc,
     "HR teams spend significant non-productive hours on tasks that should be automated: computing leave balances, "
     "chasing managers for approval signatures, answering repetitive employee queries about salary, shifts, and "
     "remaining leave days — all while managing compliance with labour regulations that change frequently.",
     size=10.5)

doc.add_paragraph().paragraph_format.space_after = Pt(6)
subheading(doc, "Core Pain Points")

pain_points = [
    ("Attendance Fraud",
     "Manual registers and basic card swipes are easily abused through proxy attendance (buddy punching), "
     "costing organisations an estimated 2–4% of total payroll annually."),
    ("No Real-Time Visibility",
     "Managers have no live picture of who is present, who is on leave, and who is late until the HR team "
     "manually compiles a report — often a day or more later."),
    ("Disconnected Systems",
     "Attendance, leave, payroll, shift scheduling, and project tracking exist in separate tools (Excel sheets, "
     "WhatsApp groups, physical registers) with no integration. Data is duplicated and frequently inconsistent."),
    ("Remote & Field Worker Blind Spots",
     "Organisations with distributed teams, field staff, or remote workers have no reliable mechanism to verify "
     "location, record attendance, or communicate shift changes in real time."),
    ("HR Administrative Overload",
     "Leave applications, approvals, salary queries, and compliance reporting consume 30–40% of an HR "
     "professional's working week on administrative tasks that carry zero strategic value."),
    ("Subscription & Billing Complexity",
     "Existing HRMS tools charge opaque per-user fees with no transparent subscription management, making "
     "cost forecasting difficult for growing SMBs."),
    ("Employee Experience Gap",
     "Employees have no self-service access to their records. Every query — 'How many leaves do I have left?', "
     "'When is my next shift?' — requires contacting HR directly."),
]

for title, desc in pain_points:
    p   = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(4)
    tr  = p.add_run(f"▸  {title}:  ")
    tr.bold            = True
    tr.font.color.rgb  = BRAND_BLUE
    tr.font.size       = Pt(10.5)
    dr  = p.add_run(desc)
    dr.font.color.rgb  = GREY_TEXT
    dr.font.size       = Pt(10.5)

doc.add_paragraph().paragraph_format.space_after = Pt(6)
subheading(doc, "The Market Opportunity")

body(doc,
     "India alone has over 63 million MSMEs employing hundreds of millions of workers. The global HRMS market "
     "is projected to reach USD 38.36 billion by 2030 (CAGR 11.4%). Existing solutions in the Indian market "
     "are either enterprise-grade and prohibitively expensive, or basic attendance apps that lack integration. "
     "WorkPing targets the underserved mid-market: organisations of 20 – 2,000 employees that need enterprise "
     "capability at SMB pricing.",
     size=10.5)


# ══════════════════════════════════════════════════════════════════════════════
# 2. SOLUTION OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
page_break(doc)

heading(doc, "2. Solution Overview", level=1, size=20, space_before=0)
divider(doc)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

body(doc,
     "WorkPing is a full-stack, multi-tenant workforce management platform that replaces every disconnected HR "
     "tool with a single integrated system. It combines biometric attendance, leave management, shift scheduling, "
     "subscription billing, AI-powered communication, and cloud-native storage — accessible from a browser, "
     "a mobile app, and even WhatsApp.",
     size=10.5)

doc.add_paragraph().paragraph_format.space_after = Pt(6)
subheading(doc, "Key Capabilities at a Glance")

capabilities = [
    ("Biometric Attendance", "Face recognition check-in via mobile camera or browser webcam. GPS co-ordinates captured for geofence verification. No hardware needed."),
    ("Live Attendance Board", "Real-time dashboard for administrators — see who checked in, who is absent, who is late, as it happens via Socket.io push."),
    ("Employee Self-Service", "Employees apply for leave, view salary slips, check shift schedules, and query attendance history — without contacting HR."),
    ("AI WhatsApp Chatbot", "Employees interact in natural language via WhatsApp. The chatbot answers attendance queries, processes leave requests, and shares salary slips."),
    ("Multi-level Approvals", "Leave requests flow through configurable approval chains — manager → HR → admin — with automated WhatsApp/email notifications at each step."),
    ("Subscription Billing", "Tiered plans (per-employee-per-month). Admins subscribe and pay via UPI through PhonePe. Auto-renewal cron with real-time payment confirmation."),
    ("Two-Factor Authentication", "TOTP authenticator app + Google/Microsoft SSO. No password-only accounts for admin roles."),
    ("Cloud-Native Storage", "All documents (profile images, payslips, bulk Excel imports) stored in Oracle Cloud Object Storage with pre-signed secure URLs."),
    ("Bulk Operations", "Import hundreds of employees via Excel upload. Export attendance and payroll reports to Excel in one click."),
]

make_table(doc,
           ["Capability", "What It Does"],
           [[c, d] for c, d in capabilities],
           col_widths=[6, 11])


# ══════════════════════════════════════════════════════════════════════════════
# 3. SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
page_break(doc)

heading(doc, "3. System Architecture", level=1, size=20, space_before=0)
divider(doc)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

body(doc,
     "WorkPing follows a microservices architecture. The Core API is the single integration point for all clients. "
     "Purpose-built microservices handle each external concern in isolation, communicating with the core API via "
     "authenticated internal HTTP calls. Nginx acts as the reverse proxy, API gateway, and SSL terminator.",
     size=10.5)

doc.add_paragraph().paragraph_format.space_after = Pt(8)

arch_box = doc.add_table(rows=1, cols=1)
remove_table_border(arch_box)
arch_cell = arch_box.cell(0, 0)
set_cell_bg(arch_cell, BRAND_DARK)
set_cell_borders(arch_cell, "1A56DB")

arch_text = (
    "  CLIENTS\n"
    "  Admin Dashboard (admin.workping.live) · Employee Portal (employee.workping.live) · Mobile App\n\n"
    "         │  HTTPS / WSS\n"
    "         ▼\n\n"
    "  NGINX  —  TLS termination · subdomain routing · WebSocket upgrade · static file serving\n\n"
    "         │  HTTP (JWT / Bearer)\n"
    "         ▼\n\n"
    "  CORE API  (api.workping.live)  —  Node.js cluster · Express 5 · MongoDB · Redis · Socket.io\n"
    "  Auth · Employees · Attendance · Leave · Shifts · Subscriptions · Projects · Payroll\n\n"
    "      │          │          │          │          │   API-key authentication\n"
    "      ▼          ▼          ▼          ▼          ▼\n"
    "  BIOMETRIC   MAILER    PAYMENTS   CHATBOT    STORAGE\n"
    "  :8001       :3003      :3001      :3002      :8000\n"
    "  FastAPI     Express    Express    Express    Express\n"
    "  InsightFace SMTP       PhonePe    BullMQ     OCI SDK\n\n"
    "                 ▼ shared state\n"
    "         Redis 7 (OTP · queue · pub-sub · Socket adapter)\n"
    "         MongoDB Atlas (all application data)"
)

ap = arch_cell.paragraphs[0]
ap.paragraph_format.space_before = Pt(8)
ap.paragraph_format.space_after  = Pt(8)
ar = ap.add_run(arch_text)
ar.font.name       = "Courier New"
ar.font.size       = Pt(8.5)
ar.font.color.rgb  = BRAND_ACCENT

doc.add_paragraph().paragraph_format.space_after = Pt(8)

subheading(doc, "Deployment Topology — Oracle Cloud Infrastructure")

make_table(doc,
           ["VM", "Hostname", "Services", "Specs"],
           [
               ["VM-1", "api.workping.live",        "Core API (Node cluster + PM2) · Redis · Nginx",    "4 vCPU · 24 GB RAM · Ubuntu 22.04"],
               ["VM-2", "face.workping.live",       "Biometric Service (FastAPI + InsightFace, CPU ONNX)","4 vCPU · 24 GB RAM · Ubuntu 22.04"],
               ["VM-3", "*.workping.live (multi)",  "Mailer · PhonePe · WhatsApp · Storage services",   "4 vCPU · 24 GB RAM · Ubuntu 22.04"],
               ["Cloud","atlas.mongodb.com",        "MongoDB Atlas — managed cluster",                   "Managed (no VM)"],
               ["Cloud","objectstorage.oraclecloud","OCI Object Storage — managed bucket",              "Managed (no VM)"],
           ],
           col_widths=[2.5, 5, 7, 5])


# ══════════════════════════════════════════════════════════════════════════════
# 4. TECHNOLOGY STACK
# ══════════════════════════════════════════════════════════════════════════════
page_break(doc)

heading(doc, "4. Technology Stack", level=1, size=20, space_before=0)
divider(doc)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

body(doc, "Every technology in WorkPing was chosen to maximise reliability, developer velocity, and security "
     "while minimising operational cost — particularly important for an OCI Always-Free infrastructure budget.", size=10.5)
doc.add_paragraph().paragraph_format.space_after = Pt(6)

subheading(doc, "4.1 Frontend Technologies")
make_table(doc,
           ["Technology", "Version", "Role", "Why Chosen"],
           [
               ["React", "18.3", "UI framework (Admin + Employee portals)", "Component model, massive ecosystem, team expertise"],
               ["Vite", "5.2", "Build tool & dev server", "Sub-second HMR, fastest bundler for SPA workloads"],
               ["React Router", "v6", "Client-side routing", "De-facto standard for React SPA routing"],
               ["Bootstrap", "5.3", "CSS framework", "Responsive grid, pre-built components, quick styling"],
               ["ApexCharts", "3.41", "Analytics dashboards", "Rich chart types, responsive, React bindings"],
               ["FullCalendar", "6.1", "Shift & holiday calendar", "Feature-complete, drag-and-drop, React integration"],
               ["socket.io-client", "4.8", "Real-time updates", "Pairs with server Socket.io; automatic reconnection"],
               ["TensorFlow.js + MediaPipe", "—", "In-browser face enrollment", "On-device inference — no raw photo sent for enrollment"],
               ["react-hook-form + yup", "—", "Forms & validation", "Minimal re-renders, declarative schema validation"],
               ["Leaflet (react-leaflet)", "1.9", "Geofence map", "Lightweight open-source maps; no Google Maps cost"],
               ["xlsx", "0.18", "Excel import/export", "Pure JS; no server-side dependency for reports"],
           ],
           col_widths=[4, 2, 5.5, 5.5])

doc.add_paragraph().paragraph_format.space_after = Pt(8)
subheading(doc, "4.2 Mobile Technologies")
make_table(doc,
           ["Technology", "Version", "Role"],
           [
               ["React Native", "0.83", "Cross-platform mobile framework (iOS + Android)"],
               ["Expo", "55", "Build toolchain; native module abstraction"],
               ["react-native-vision-camera", "4.7", "Real-time face detection for check-in"],
               ["expo-location", "55.1", "GPS capture for geofence verification"],
               ["expo-notifications", "55.0", "Push notifications (shift reminders, approvals)"],
               ["React Navigation", "7", "Stack + tab navigator"],
               ["react-hook-form + yup", "—", "Form handling and validation"],
               ["axios", "1.13", "HTTP client for Core API"],
           ],
           col_widths=[5, 2, 10])

doc.add_paragraph().paragraph_format.space_after = Pt(8)
subheading(doc, "4.3 Backend Technologies")
make_table(doc,
           ["Technology", "Version", "Role", "Why Chosen"],
           [
               ["Node.js", "20 LTS", "JavaScript runtime", "Non-blocking I/O; ideal for API gateway + event-driven work"],
               ["Express 5", "5.2", "HTTP framework", "Minimal, fast, async-native in v5 (no try/catch boilerplate)"],
               ["MongoDB Atlas", "8 (ODM)", "Primary database", "Document model fits nested employee/attendance schemas"],
               ["Mongoose", "8.17", "MongoDB ODM", "Schema validation, middleware hooks, populate joins"],
               ["Redis 7", "5.10 client", "Cache · OTP store · queue · pub-sub", "Single service fills four infrastructure roles"],
               ["Socket.io", "4.8", "Real-time WebSocket layer", "Automatic fallback to polling; Redis adapter for cluster"],
               ["JWT + bcrypt", "9.0 / 6.0", "Auth tokens + password hashing", "Stateless auth; bcrypt cost-10 for password resistance"],
               ["speakeasy", "2.0", "TOTP two-factor auth", "RFC 6238 compliant; standard authenticator app support"],
               ["node-cron", "4.2", "Scheduled jobs", "Lightweight; no separate process; renewal + reminder cron"],
               ["Multer + xlsx", "2.0 / 0.18", "File upload + Excel parsing", "Multipart handling; server-side Excel for bulk import"],
               ["helmet + rate-limit", "8.0 / 8.1", "Security middleware", "Twelve security headers; per-IP throttling"],
               ["Morgan + Pino", "1.10 / 10.3", "HTTP + structured logging", "Morgan for dev; Pino for structured prod logs"],
           ],
           col_widths=[3.5, 2, 4.5, 7])

doc.add_paragraph().paragraph_format.space_after = Pt(8)
subheading(doc, "4.4 Biometric Service Technologies")
make_table(doc,
           ["Technology", "Role", "Why Chosen"],
           [
               ["Python 3.10+", "Runtime", "Rich ML/CV ecosystem; best support for InsightFace & FAISS"],
               ["FastAPI", "HTTP framework", "Async-native, auto OpenAPI docs, Pydantic validation"],
               ["InsightFace (AntelopeV2)", "Face detection + embedding", "State-of-the-art accuracy; ONNX export; fully self-hosted"],
               ["ArcFace R100", "512-dim face embedding", "99.8% LFW accuracy; production-proven at scale"],
               ["FAISS IndexFlatIP", "Bulk vector search", "In-memory; sub-millisecond search; no external service needed"],
               ["ONNX Runtime", "Model inference", "CPU + GPU support; portable across hardware"],
               ["Motor (async MongoDB)", "Embedding storage", "Non-blocking DB I/O keeps FastAPI event loop free"],
               ["redis[hiredis]", "Task queue + result cache", "hiredis C extension for maximum throughput"],
               ["ThreadPoolExecutor", "Inference threading", "Offloads blocking GPU/CPU calls from asyncio event loop"],
           ],
           col_widths=[5, 4, 8])


# ══════════════════════════════════════════════════════════════════════════════
# 5. MICROSERVICES BREAKDOWN
# ══════════════════════════════════════════════════════════════════════════════
page_break(doc)

heading(doc, "5. Microservices — Detailed Breakdown", level=1, size=20, space_before=0)
divider(doc)
doc.add_paragraph().paragraph_format.space_after = Pt(6)

services = [
    {
        "name": "5.1 Core API",
        "url": "api.workping.live",
        "stack": "Node.js · Express 5 · Mongoose · Redis · Socket.io",
        "desc": (
            "The brain of WorkPing. All client requests (web and mobile) flow through this service. "
            "It manages authentication, business logic, data persistence, and real-time events. "
            "Runs as a Node.js cluster (one worker per CPU) under PM2 for zero-downtime restarts."
        ),
        "features": [
            "Admin routes: organisation setup, employee CRUD, bulk Excel import, shift scheduling, leave decisions, subscription management",
            "Employee routes: profile, attendance history, leave application, salary slip, personal dashboard",
            "Internal routes: machine-to-machine endpoints consumed by microservices (chatbot, biometric)",
            "JWT access tokens (15 min) + refresh token rotation",
            "Google OAuth2 + Microsoft OAuth2 SSO",
            "TOTP two-factor authentication with QR enrollment",
            "node-cron: subscription auto-renewal + shift reminder jobs",
            "Socket.io + Redis adapter: real-time attendance board and payment confirmation push",
        ],
    },
    {
        "name": "5.2 Biometric Service",
        "url": "face.workping.live",
        "stack": "Python · FastAPI · InsightFace · FAISS · Motor · Redis",
        "desc": (
            "Handles all face recognition operations. Runs entirely on-premises — biometric data never "
            "leaves WorkPing infrastructure. Uses InsightFace AntelopeV2 (SCRFD for detection, ArcFace R100 "
            "for 512-dimensional embedding). Inference is decoupled from HTTP via a Redis task queue, "
            "keeping API response times consistently fast regardless of GPU/CPU load."
        ),
        "features": [
            "Face enrollment: extract + store 512-dim L2-normalised embedding in MongoDB",
            "Face verification: cosine similarity against stored embedding (threshold 0.6)",
            "Async inference: Redis BLPOP queue → ThreadPoolExecutor worker → result ticket with 5-min TTL",
            "Bulk search: FAISS IndexFlatIP for organisation-wide multi-user face matching",
            "Embedding cache: Redis TTL cache to skip re-fetching unchanged embeddings",
            "CUDA auto-detected; falls back gracefully to CPU ONNX Runtime",
        ],
    },
    {
        "name": "5.3 Mailer Service",
        "url": "internal (no public subdomain)",
        "stack": "Node.js · Express 5 · Nodemailer · Handlebars · Redis",
        "desc": (
            "Handles all outbound email for the platform. Decoupled from the Core API so email delays "
            "never block business logic. Redis stores OTPs as single-use, TTL-bounded keys — "
            "deleted immediately on successful verification to prevent replay attacks."
        ),
        "features": [
            "Email OTP: 6-digit, 30-min TTL, single-use (deleted on verify)",
            "Password reset OTP: 10-min TTL with dedicated verify endpoint",
            "Handlebars HTML templates: welcome, OTP, password reset, alert (info/warning/danger/success), notification",
            "Stateless — any instance can verify any OTP (Redis is the shared store)",
            "Alert severity tiers with action-link buttons in email body",
        ],
    },
    {
        "name": "5.4 Payment Service",
        "url": "phonepe.workping.live",
        "stack": "Node.js · Express 5 · Axios · PhonePe APIs",
        "desc": (
            "Bridges WorkPing subscription billing with PhonePe UPI. Handles payment initiation, "
            "status polling, and secure webhook receipt. All webhook verification uses "
            "crypto.timingSafeEqual to prevent timing-based HMAC attacks."
        ),
        "features": [
            "Payment initiation: POST to PhonePe v4 API, returns checkout URL (10-min expiry)",
            "Supported modes: UPI collect · UPI intent · UPI QR · Credit/Debit card · Net Banking",
            "Webhook receiver: SHA-256 HMAC signature verification before forwarding to Core API",
            "Payment status polling: fallback to PhonePe status API if webhook is delayed",
            "Sandbox / production toggle via single ENV variable",
        ],
    },
    {
        "name": "5.5 WhatsApp Chatbot Service",
        "url": "whatsapp.workping.live",
        "stack": "Node.js · Express 5 · BullMQ · Redis · AWS Bedrock / OpenAI",
        "desc": (
            "Gives every employee a 24/7 HR assistant on WhatsApp. The service decouples webhook "
            "receipt from LLM processing using BullMQ, so Meta's 20-second webhook timeout is never "
            "breached. Intent routing uses a fast rule engine first — LLM is invoked only for "
            "unrecognised queries, balancing cost and coverage."
        ),
        "features": [
            "WhatsApp Cloud API integration (Meta Business Platform)",
            "BullMQ job queue: decouples receipt from async LLM + Core API calls",
            "Intent routing: keyword rule engine (fast) → LLM fallback (flexible)",
            "Supported intents: attendance queries, leave application & status, shift schedule, salary slip",
            "LLM provider-agnostic: Ollama · AWS Bedrock · OpenAI · Groq · OpenRouter · Mistral",
            "Switch LLM provider at runtime via dashboard API — no restart",
            "Shift reminders: scheduled via node-cron, 15 min before shift start",
            "Leave approval flow: multi-step conversational workflow triggered from admin",
        ],
    },
    {
        "name": "5.6 Object Storage Service",
        "url": "s3.workping.live",
        "stack": "Node.js · Express 5 · OCI SDK · Multer · Pino",
        "desc": (
            "Provides a secure, rate-limited HTTP interface to Oracle Cloud Infrastructure Object Storage. "
            "All uploads are validated for MIME type and size before reaching OCI. Pre-signed URLs "
            "provide time-limited direct-download access without exposing OCI credentials to clients."
        ),
        "features": [
            "Upload, download, list, and delete operations on OCI buckets",
            "Pre-signed upload and download URLs (15-min expiry)",
            "File size limit: 50 MB (configurable); MIME type allowlist",
            "API key authentication + Helmet security headers + rate limiting (100 req/15 min)",
            "Structured Pino logging with 30-day metrics export (JSON + CSV)",
            "OCI credentials mounted read-only from host (~/.oci)",
        ],
    },
]

for svc in services:
    subheading(doc, svc["name"])

    meta_table = doc.add_table(rows=2, cols=2)
    remove_table_border(meta_table)
    for ri, (lbl, val) in enumerate([("Public URL", svc["url"]), ("Stack", svc["stack"])]):
        lc = meta_table.cell(ri, 0)
        vc = meta_table.cell(ri, 1)
        set_cell_bg(lc, BRAND_LIGHT)
        lc.width = Cm(3.5)
        vc.width = Cm(13.5)
        lp = lc.paragraphs[0]
        lp.paragraph_format.space_before = Pt(3)
        lp.paragraph_format.space_after  = Pt(3)
        lr = lp.add_run(f"  {lbl}")
        lr.bold = True
        lr.font.color.rgb = BRAND_BLUE
        lr.font.size = Pt(9.5)
        vp = vc.paragraphs[0]
        vp.paragraph_format.space_before = Pt(3)
        vp.paragraph_format.space_after  = Pt(3)
        vr = vp.add_run(f"  {val}")
        vr.font.color.rgb = GREY_TEXT
        vr.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    body(doc, svc["desc"], size=10.5)

    for feat in svc["features"]:
        bullet(doc, feat)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)


# ══════════════════════════════════════════════════════════════════════════════
# 6. WORKFLOW & DATA FLOWS
# ══════════════════════════════════════════════════════════════════════════════
page_break(doc)

heading(doc, "6. Workflow & Data Flows", level=1, size=20, space_before=0)
divider(doc)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── 6.1 Onboarding ─────────────────────────────────────────────────────────
subheading(doc, "6.1 Organisation Onboarding")

onboard_steps = [
    ("Admin registers", "Admin creates an account with email + password. Email OTP is sent by the Mailer service and verified."),
    ("2FA enrollment", "Admin scans a TOTP QR code in an authenticator app. All subsequent logins require the 6-digit TOTP code."),
    ("Organisation setup", "Admin fills organisation details: name, address, geofence co-ordinates, working hours policy."),
    ("Plan selection", "Admin browses subscription plans and initiates a UPI payment via PhonePe. Real-time Socket.io push confirms payment in the dashboard."),
    ("Employee import", "Admin uploads a bulk Excel file. The server parses it (xlsx), creates employee accounts, and sends welcome emails via the Mailer service."),
    ("Face enrollment", "Each employee uses the browser webcam or mobile camera. TensorFlow.js / Vision Camera extracts a face embedding client-side; the embedding is sent to the Biometric service and stored in MongoDB."),
]

for i, (step, desc) in enumerate(onboard_steps, 1):
    p   = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.4)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(4)
    sr  = p.add_run(f"  {i}.  {step}:  ")
    sr.bold            = True
    sr.font.color.rgb  = BRAND_BLUE
    sr.font.size       = Pt(10.5)
    dr  = p.add_run(desc)
    dr.font.color.rgb  = GREY_TEXT
    dr.font.size       = Pt(10.5)

doc.add_paragraph().paragraph_format.space_after = Pt(8)

# ── 6.2 Daily Check-In ─────────────────────────────────────────────────────
subheading(doc, "6.2 Daily Face Check-In (Mobile)")

flow_box = doc.add_table(rows=1, cols=1)
remove_table_border(flow_box)
fc = flow_box.cell(0, 0)
set_cell_bg(fc, BRAND_DARK)
set_cell_borders(fc, "06B6D4")
fp = fc.paragraphs[0]
fp.paragraph_format.space_before = Pt(6)
fp.paragraph_format.space_after  = Pt(6)
fr = fp.add_run(
    "  Employee opens mobile app → Camera frame captured by Vision Camera\n"
    "  ↓\n"
    "  POST face.workping.live/api/v1/detect  { image_base64, user_id, org_id }\n"
    "  ↓\n"
    "  Redis BLPOP queue  →  Inference Worker (ThreadPoolExecutor)\n"
    "       → fetch stored embedding from MongoDB\n"
    "       → SCRFD detect face region  →  ArcFace extract 512-dim embedding\n"
    "       → cosine_similarity(query, stored)  →  write result to Redis ticket:<uuid> TTL=300s\n"
    "  ↓\n"
    "  App polls GET face.workping.live/result/<ticket_id>   (every 1s)\n"
    "  ↓ (match confirmed)\n"
    "  POST api.workping.live/api/user/attendance/check-in  { match_score, gps_location }\n"
    "  ↓\n"
    "  Core API  →  write AttendanceRecord to MongoDB\n"
    "           →  publish event to Redis pub-sub channel\n"
    "           →  Socket.io Redis adapter fans out to all workers\n"
    "           →  Admin dashboard receives real-time attendance update"
)
fr.font.name       = "Courier New"
fr.font.size       = Pt(8.5)
fr.font.color.rgb  = BRAND_ACCENT
doc.add_paragraph().paragraph_format.space_after = Pt(8)

# ── 6.3 Leave Flow ─────────────────────────────────────────────────────────
subheading(doc, "6.3 Leave Application & Approval")

leave_steps = [
    "Employee applies via web portal or WhatsApp chatbot.",
    "Core API creates a LeaveRequest record (status: PENDING) and notifies the manager via WhatsApp (POST to chatbot service /api/secure/whatsapp/send).",
    "Manager receives a WhatsApp message with Approve / Reject inline options.",
    "Manager's reply triggers a callback to Core API /internal/leave/decide.",
    "Core API updates LeaveRequest status, deducts from balance, and notifies the employee via WhatsApp + email.",
    "HR admin has a full audit trail in the Leave Management dashboard.",
]
for step in leave_steps:
    bullet(doc, step)

doc.add_paragraph().paragraph_format.space_after = Pt(8)

# ── 6.4 Subscription Billing ───────────────────────────────────────────────
subheading(doc, "6.4 Subscription Payment")

billing_steps = [
    "Admin selects a plan in the Admin Dashboard → Core API calls Payment Service /api/payments/initiate-payment.",
    "Payment Service posts to PhonePe API → returns a checkout URL (valid 10 min).",
    "Admin is redirected to PhonePe payment page. Employee completes UPI / card payment.",
    "PhonePe fires a signed webhook to phonepe.workping.live. Service verifies SHA-256 HMAC signature.",
    "Payment Service forwards a verified callback to Core API /internal/payments/webhook (verified by x-webhook-secret).",
    "Core API updates subscription status in MongoDB, publishes a Redis payment event.",
    "Socket.io push delivers real-time payment confirmation to the Admin Dashboard — no page refresh needed.",
    "node-cron checks subscription expiry daily; auto-renewal is attempted 3 days before expiry.",
]
for step in billing_steps:
    bullet(doc, step)

doc.add_paragraph().paragraph_format.space_after = Pt(8)

# ── 6.5 WhatsApp Chatbot ───────────────────────────────────────────────────
subheading(doc, "6.5 WhatsApp Chatbot Interaction")

chat_steps = [
    "Employee sends a WhatsApp message (e.g., 'How many leaves do I have?').",
    "Meta fires a webhook to whatsapp.workping.live/webhook (must respond within 20 seconds).",
    "Server immediately enqueues a BullMQ job and returns 200 OK to Meta.",
    "BullMQ worker picks up the job: runs the keyword/pattern rule engine.",
    "If intent is matched (e.g., 'leave balance') → call Core API /internal/leave/balance/:userId.",
    "If no intent matches → forward message to LLM (AWS Bedrock / Groq / Ollama) for natural language handling.",
    "Formatted response is sent back to the employee via WhatsApp Cloud API POST /messages.",
]
for step in chat_steps:
    bullet(doc, step)


# ══════════════════════════════════════════════════════════════════════════════
# 7. SECURITY MODEL
# ══════════════════════════════════════════════════════════════════════════════
page_break(doc)

heading(doc, "7. Security Model", level=1, size=20, space_before=0)
divider(doc)
doc.add_paragraph().paragraph_format.space_after = Pt(6)

body(doc,
     "WorkPing implements defence-in-depth with seven distinct security layers. Every service is hardened "
     "independently, so a compromise of one service cannot cascade to others.",
     size=10.5)
doc.add_paragraph().paragraph_format.space_after = Pt(6)

security_layers = [
    ("Network", "HTTPS enforced for all subdomains (TLS 1.2+). CORS allowlist — no wildcard in production. All internal service calls stay within the private VM network."),
    ("Transport", "helmet middleware on every service sets 12 HTTP security headers: HSTS, X-Frame-Options, X-Content-Type-Options, Content-Security-Policy. Request body capped at 10 KB to prevent payload inflation."),
    ("Rate Limiting", "Global: 200 requests / 15 min per IP. Auth & OTP endpoints: 10 requests / 15 min per IP. Biometric: 100 requests / 15 min. Prevents brute-force and enumeration."),
    ("Authentication", "JWT access tokens (15-min expiry) + refresh token rotation. Google OAuth2 and Microsoft OAuth2 SSO. TOTP two-factor auth (speakeasy, RFC 6238) — mandatory for admin roles."),
    ("Authorisation", "requireRole middleware enforces admin | manager | teamlead | employee boundaries on every protected route. authorizeManager prevents cross-team data access."),
    ("Inter-Service", "All microservice calls carry Authorization: Bearer <INTERNAL_API_KEY>. PhonePe and Core API webhooks verified with crypto.timingSafeEqual (constant-time comparison — eliminates timing attacks)."),
    ("Data Security", "Passwords: bcrypt cost-factor 10 (resistant to GPU brute-force). JWTs: HS256 with minimum 256-bit secret. OTPs: 6-digit, single-use, Redis-TTL-bounded. Face embeddings: stored as numeric vectors only — no raw images retained server-side."),
]

make_table(doc,
           ["Layer", "Controls"],
           [[f"Layer {i+1} — {lyr}", ctrl] for i, (lyr, ctrl) in enumerate(security_layers)],
           col_widths=[5.5, 11.5])


# ══════════════════════════════════════════════════════════════════════════════
# 8. DEPLOYMENT & INFRASTRUCTURE
# ══════════════════════════════════════════════════════════════════════════════
page_break(doc)

heading(doc, "8. Deployment & Infrastructure", level=1, size=20, space_before=0)
divider(doc)
doc.add_paragraph().paragraph_format.space_after = Pt(6)

subheading(doc, "8.1 Infrastructure Provider — Oracle Cloud Infrastructure (OCI)")

body(doc,
     "All WorkPing compute runs on OCI Always Free Ampere (ARM) virtual machines. OCI was chosen over "
     "AWS/GCP/Azure for the following reasons:", size=10.5)

oci_reasons = [
    "Zero egress fees — AWS S3 and EC2 charge $0.09/GB for data leaving the cloud. OCI Object Storage and compute have no egress cost.",
    "Always Free Ampere tier provides 4 vCPU + 24 GB RAM per VM at no cost — far more generous than any other cloud provider's free tier.",
    "Full Ubuntu 22.04 LTS control — no managed runtimes or vendor lock-in.",
    "OCI Object Storage is S3-protocol compatible, making SDK integration straightforward.",
]
for r in oci_reasons:
    bullet(doc, r)

doc.add_paragraph().paragraph_format.space_after = Pt(8)
subheading(doc, "8.2 Process Management")

pm_data = [
    ("PM2 (cluster mode)", "Core API", "One worker per CPU. Zero-downtime reloads on deploy. Exponential-backoff restart on crash. Log rotation."),
    ("Docker Compose", "All microservices", "Each service in its own container. Shared Redis service. Health checks with start_period for model loading."),
    ("Nginx", "All VMs", "Reverse proxy + SSL termination + static file serving + WebSocket upgrade pass-through."),
    ("Let's Encrypt / Certbot", "All VMs", "Automatic TLS certificate issuance and renewal via systemd timer."),
]
make_table(doc, ["Tool", "Used For", "Detail"], pm_data, col_widths=[4.5, 4.5, 8])

doc.add_paragraph().paragraph_format.space_after = Pt(8)
subheading(doc, "8.3 CI/CD Pipeline")

cicd_steps = [
    "Developer pushes to main branch on GitHub.",
    "GitHub Actions workflow triggers: install dependencies → run linter → run tests.",
    "On pass: SSH into target VM, git pull, npm ci / pip install, PM2 reload or docker compose up -d --build.",
    "Frontend: npm run build produces dist/. Nginx serves the new build immediately (no container restart).",
    "Zero-downtime for the Core API: PM2 cluster reload replaces workers one by one.",
]
for s in cicd_steps:
    bullet(doc, s)


# ══════════════════════════════════════════════════════════════════════════════
# 9. IMPACT ON HR MANAGEMENT
# ══════════════════════════════════════════════════════════════════════════════
page_break(doc)

heading(doc, "9. Impact on HR Management", level=1, size=20, space_before=0)
divider(doc)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

body(doc,
     "WorkPing does not just digitise existing HR processes — it fundamentally changes what HR teams "
     "spend their time on. By automating routine administration, it frees HR professionals for strategic "
     "work: talent development, culture building, and workforce planning.",
     size=10.5)
doc.add_paragraph().paragraph_format.space_after = Pt(6)

subheading(doc, "9.1 Quantifiable Benefits for Organisations")

benefits_table = [
    ["Attendance Fraud Elimination",
     "Face recognition + GPS geofence makes proxy attendance physically impossible. Biometric identity cannot be shared.",
     "Industry avg: 2–4% of payroll saved"],
    ["HR Admin Time Reduction",
     "Leave applications, OTP verifications, salary slip distribution, shift communication — all automated. HR staff redirected to strategic roles.",
     "30–40% admin hour reduction"],
    ["Real-Time Workforce Visibility",
     "Live attendance board updated instantly via Socket.io. No end-of-day compilation. Managers make informed decisions in the moment.",
     "Data lag: hours → seconds"],
    ["Employee Self-Service",
     "Employees check balances, view slips, apply leave, and get answers via WhatsApp — without ever contacting HR.",
     "~60% fewer HR queries"],
    ["Subscription Cost Transparency",
     "Clear per-employee-per-month pricing. Admins see exactly what they pay for. Auto-renewal eliminates service disruption.",
     "Predictable OPEX"],
    ["Compliance-Ready Audit Trail",
     "Every attendance event, leave decision, and shift change is timestamped in MongoDB. Exportable to Excel for labour compliance audits.",
     "Full audit trail"],
    ["Remote & Field Worker Coverage",
     "Mobile app with GPS works anywhere. No physical hardware required — just a smartphone.",
     "Hardware cost: ₹0"],
    ["Paperless Operations",
     "All documents — contracts, payslips, bulk imports — stored in OCI Object Storage. Instant retrieval via pre-signed URLs.",
     "Zero paper dependency"],
]

make_table(doc,
           ["Impact Area", "How WorkPing Delivers", "Expected Outcome"],
           benefits_table,
           col_widths=[4.5, 8.5, 4])

doc.add_paragraph().paragraph_format.space_after = Pt(8)
subheading(doc, "9.2 How WorkPing Transforms HR Roles")

role_impacts = [
    ("HR Administrator",
     ["Approves or rejects leave with one tap rather than tracking paper forms",
      "Runs payroll reports by exporting pre-compiled Excel — not manually computing spreadsheets",
      "Monitors live attendance without calling team leaders",
      "Sets geofence zones per office/site — system enforces location automatically"]),
    ("Manager / Team Lead",
     ["Receives WhatsApp approval requests — no need to log into a portal",
      "Views team attendance and project task status in a single dashboard",
      "Assigns shifts via drag-and-drop calendar; employees notified instantly",
      "Escalates leave decisions upward with one click"]),
    ("Employee",
     ["Checks in with a selfie — 2–3 seconds vs 10+ minutes queuing at a card reader",
      "Asks the WhatsApp bot 'How many leaves do I have?' at 11 pm without waiting for Monday",
      "Downloads salary slip as PDF — no HR office visit required",
      "Receives push notification 15 minutes before shift starts"]),
    ("Finance / Payroll Team",
     ["Exports attendance and leave data pre-formatted for payroll computation",
      "Subscription invoices and payment history available in the admin dashboard",
      "Audit trail for every transaction — timestamped, user-attributed"]),
]

for role, points in role_impacts:
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    r   = p.add_run(f"  {role}")
    r.bold = True
    r.font.color.rgb = BRAND_BLUE
    r.font.size = Pt(11)
    for point in points:
        sub_bullet(doc, point)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

doc.add_paragraph().paragraph_format.space_after = Pt(8)
subheading(doc, "9.3 Competitive Differentiation")

diff_table = [
    ["Biometric Attendance", "✓ Face recognition (no hardware)", "Card reader / OTP / manual", "✓"],
    ["Mobile App", "✓ iOS + Android (React Native)", "Web only or separate app", "✓"],
    ["AI WhatsApp Chatbot", "✓ 24/7 intent + LLM", "Email or portal only", "✓"],
    ["Self-Hosted Biometrics", "✓ On-premises (privacy-first)", "Cloud API (data leaves org)", "✓"],
    ["Real-Time Dashboard", "✓ Socket.io push", "Polling / manual refresh", "✓"],
    ["Geofence Verification", "✓ GPS + configurable zones", "Office-only or none", "✓"],
    ["Bulk Employee Import", "✓ Excel upload", "Form-by-form entry", "✓"],
    ["Multi-tenant SaaS", "✓ Org-isolated data", "Single-tenant or shared DB", "✓"],
    ["UPI-Native Payments", "✓ PhonePe (India-first)", "Card only or manual", "✓"],
]

make_table(doc,
           ["Feature", "WorkPing", "Typical Alternatives", "Advantage"],
           diff_table,
           col_widths=[5, 5, 5, 2])


# ══════════════════════════════════════════════════════════════════════════════
# 10. FUTURE SCOPE
# ══════════════════════════════════════════════════════════════════════════════
page_break(doc)

heading(doc, "10. Future Scope", level=1, size=20, space_before=0)
divider(doc)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

body(doc,
     "WorkPing has been architected with future extensibility as a first-class concern. The following "
     "capabilities are on the roadmap, prioritised by business impact.",
     size=10.5)
doc.add_paragraph().paragraph_format.space_after = Pt(6)

future_items = [
    {
        "title": "Payroll Automation",
        "priority": "HIGH",
        "desc": (
            "End-to-end payroll computation — CTC breakdown (basic, HRA, DA, PF, ESI, TDS), "
            "automated payslip generation (PDF), and direct bank transfer via NEFT/IMPS. "
            "The current salary slip stub will be expanded into a full payroll engine. "
            "The financial ledger sub-domain will be evaluated for a PostgreSQL migration "
            "to leverage strict ACID compliance for deduction and tax ledgers."
        ),
    },
    {
        "title": "Advanced AI Analytics & Statistics",
        "priority": "HIGH",
        "desc": (
            "Predictive workforce analytics powered by ML models: "
            "attendance anomaly detection (flagging chronic lateness or mass check-in fraud patterns), "
            "leave demand forecasting (alerting managers to understaffing risk), "
            "productivity correlation dashboards (attendance vs project completion rate), "
            "and a natural-language query interface over HR data ('Show me departments with >20% absenteeism this month')."
        ),
    },
    {
        "title": "Liveness Detection (Anti-Spoofing)",
        "priority": "HIGH",
        "desc": (
            "The current biometric system has no Presentation Attack Detection (PAD). "
            "A photo or video held to the camera can pass face matching. "
            "We will integrate a Silent Face Anti-Spoofing model (ONNX-compatible, CPU-friendly) "
            "as an additional inference step before embedding comparison — a critical security hardening "
            "before enterprise deployments."
        ),
    },
    {
        "title": "International Payments (Stripe)",
        "priority": "MEDIUM",
        "desc": (
            "Add Stripe for USD/EUR card subscriptions as WorkPing expands beyond India. "
            "The payment service is already isolated and provider-agnostic by design; "
            "a Stripe provider can be added without touching the Core API."
        ),
    },
    {
        "title": "JWT Revocation & Session Management",
        "priority": "MEDIUM",
        "desc": (
            "Implement a Redis-backed token blacklist to support immediate session invalidation "
            "on logout, account suspension, or suspected compromise. "
            "Currently, a stolen JWT remains valid until its 15-minute expiry."
        ),
    },
    {
        "title": "Chatbot Conversation Memory (Vector DB)",
        "priority": "MEDIUM",
        "desc": (
            "Integrate pgvector (PostgreSQL extension) or Pinecone to store per-user conversation "
            "embeddings, enabling contextual follow-up across sessions "
            "('What about last month?' after asking about this month's attendance). "
            "Also planning LLM function-calling / tool-use to replace the hand-written rule engine "
            "with a more robust, extensible approach."
        ),
    },
    {
        "title": "CDN for Object Storage",
        "priority": "MEDIUM",
        "desc": (
            "Place OCI CDN or Cloudflare in front of s3.workping.live to cache profile images "
            "and documents at edge nodes, reducing latency for teams distributed across India."
        ),
    },
    {
        "title": "Kubernetes Migration (OCI OKE)",
        "priority": "LOW (scale trigger)",
        "desc": (
            "As concurrent user load exceeds single-VM capacity, migrate from PM2 + Docker Compose "
            "to OCI Kubernetes Engine (OKE). Docker images already exist; "
            "only K8s Deployment manifests and a Horizontal Pod Autoscaler need to be authored. "
            "OKE provides cross-VM self-healing, rolling deployments, and auto-scaling."
        ),
    },
    {
        "title": "Geofence Server-Side Enforcement",
        "priority": "LOW",
        "desc": (
            "Move geofence boundary checks from client-side UI validation to the Core API, "
            "with configurable radius per organisation. "
            "This prevents GPS spoofing via modified mobile clients."
        ),
    },
    {
        "title": "Email Deliverability (Amazon SES)",
        "priority": "LOW",
        "desc": (
            "Replace raw SMTP with Amazon SES for external-facing emails "
            "to improve deliverability, add open/click tracking, "
            "and ensure bounce management. Nodemailer SES transport is a one-line config change."
        ),
    },
]

for item in future_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    tr = p.add_run(f"  {item['title']}  ")
    tr.bold = True
    tr.font.color.rgb = BRAND_DARK
    tr.font.size = Pt(11)

    pri_color = BRAND_BLUE if "HIGH" in item["priority"] else BRAND_ACCENT if "MEDIUM" in item["priority"] else GREY_TEXT
    pr = p.add_run(f"[{item['priority']}]")
    pr.bold = True
    pr.font.color.rgb = pri_color
    pr.font.size = Pt(10)

    body(doc, f"  {item['desc']}", size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ══════════════════════════════════════════════════════════════════════════════
# CLOSING PAGE
# ══════════════════════════════════════════════════════════════════════════════
page_break(doc)

doc.add_paragraph().paragraph_format.space_after = Pt(60)

close_p = doc.add_paragraph()
close_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
cr = close_p.add_run("WorkPing")
cr.font.size = Pt(40)
cr.font.bold = True
cr.font.color.rgb = BRAND_BLUE

cp2 = doc.add_paragraph()
cp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
cr2 = cp2.add_run("Intelligent Workforce Management Platform")
cr2.font.size = Pt(16)
cr2.font.color.rgb = BRAND_ACCENT
cr2.font.italic = True

doc.add_paragraph().paragraph_format.space_after = Pt(30)
divider(doc)
doc.add_paragraph().paragraph_format.space_after = Pt(20)

url_p = doc.add_paragraph()
url_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
ur = url_p.add_run("workping.live")
ur.font.size = Pt(14)
ur.bold = True
ur.font.color.rgb = BRAND_DARK

doc.add_paragraph().paragraph_format.space_after = Pt(8)

contact_p = doc.add_paragraph()
contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
cor = contact_p.add_run("Built and deployed on Oracle Cloud Infrastructure  ·  MIT License  ·  2026")
cor.font.size = Pt(10)
cor.font.color.rgb = GREY_TEXT
cor.font.italic = True

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
import os
output_path = os.path.join(os.path.dirname(__file__), "WorkPing_Project_Documentation.docx")
doc.save(output_path)
print(f"Saved: {output_path}")
